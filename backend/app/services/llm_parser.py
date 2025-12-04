"""
LLM-powered resume parser using OpenAI GPT-4
More accurate than regex-based parsing
"""
import json
import re
from typing import Dict, Any, Optional
from pathlib import Path
from app.core.config import get_settings
from app.schemas.resume import Resume

# Optional imports
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    openai = None
    OPENAI_AVAILABLE = False

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    pypdf = None
    PYPDF_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False


class LLMResumeParser:
    """
    Use GPT-4 to extract structured resume data from text.
    Falls back to basic extraction if LLM fails.
    """

    def __init__(self):
        settings = get_settings()
        self.api_key = settings.OPENAI_API_KEY if hasattr(settings, 'OPENAI_API_KEY') else None

        if self.api_key and OPENAI_AVAILABLE:
            openai.api_key = self.api_key
        elif self.api_key and not OPENAI_AVAILABLE:
            print("Warning: OPENAI_API_KEY set but 'openai' module not installed.")

    async def parse_file(self, file_path: Path, filename: str) -> Resume:
        """Parse resume file using LLM extraction"""
        # Extract raw text
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            text = await self._extract_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            text = await self._extract_docx(file_path)
        elif extension == ".txt":
            text = await self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

        if not text or len(text.strip()) < 50:
            raise ValueError("Resume content is too short or empty")

        # Use LLM parsing for accurate extraction of all sections
        if self.api_key and OPENAI_AVAILABLE:
            try:
                return await self._parse_with_llm(text)
            except Exception as e:
                print(f"LLM parsing failed: {e}, falling back to regex parser")
                return await self._parse_basic(text)
        else:
            print("LLM parsing not available, using regex-based parser")
            return await self._parse_basic(text)

    async def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF - try PyMuPDF first for better quality"""

        # Try PyMuPDF first (better text extraction)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    # Use 'text' mode which preserves layout better for two-column resumes
                    text_parts.append(page.get_text("text"))
                doc.close()
                full_text = "\n".join(text_parts)

                # DO NOT collapse multiple spaces - they indicate two-column layouts!
                # Just return the text as-is to preserve spatial relationships
                return full_text
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}, falling back to pypdf")

        # Fallback to pypdf
        if not PYPDF_AVAILABLE:
            raise ValueError("No PDF parsing library available. Install pymupdf or pypdf.")

        text_parts = []
        with open(file_path, "rb") as file:
            reader = pypdf.PdfReader(file)
            for page in reader.pages:
                # Try layout mode first (newer pypdf versions) - preserves spacing
                try:
                    page_text = page.extract_text(extraction_mode="layout")
                except TypeError:
                    # Fallback for older pypdf versions
                    page_text = page.extract_text()

                if page_text:
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts)

        # DO NOT aggressively clean up spaces - they're important for two-column layouts!
        # Only do minimal cleanup to remove truly excessive spacing (5+ spaces)
        import re
        full_text = re.sub(r' {5,}', '    ', full_text)  # Keep 4 spaces to preserve column indicators

        return full_text

    async def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx module not installed. Cannot parse DOCX.")

        doc = Document(file_path)
        text_parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    async def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _is_bullet_point(self, line: str) -> bool:
        """
        Multi-strategy bullet detection
        Returns True if line appears to be a bullet point
        """
        line_stripped = line.strip()
        if not line_stripped:
            return False

        # Strategy 1: Bullet characters
        bullet_chars = ['•', '●', '-', '*', '○', '▪', '▫', '■', '□', '◦', '‣', '⁃', '▸', '▹', '►', '▻']
        if any(line_stripped.startswith(char) for char in bullet_chars):
            return True

        # Strategy 2: Numbered bullets (1. 2. 3.)
        if re.match(r'^\d+\.\s', line_stripped):
            return True

        # Strategy 3: Indented lines (4+ spaces or tab) - but not section headers
        if (line.startswith('    ') or line.startswith('\t')) and not self._detect_section(line):
            return True

        # Strategy 4: Action verbs (common resume action verbs)
        action_verbs = [
            'built', 'developed', 'created', 'designed', 'implemented', 'led', 'managed',
            'improved', 'increased', 'reduced', 'achieved', 'delivered', 'launched',
            'established', 'optimized', 'automated', 'collaborated', 'coordinated',
            'spearheaded', 'architected', 'engineered', 'executed', 'analyzed',
            'researched', 'integrated', 'migrated', 'deployed', 'configured',
            'maintained', 'supported', 'troubleshot', 'resolved', 'enhanced'
        ]
        first_word = line_stripped.split()[0].lower() if line_stripped.split() else ""
        if first_word in action_verbs:
            return True

        return False

    def _detect_section(self, line: str) -> Optional[str]:
        """
        Flexible section header detection
        Returns section type if line is a section header, None otherwise
        """
        line_clean = line.strip().upper()

        # Section patterns with variations
        section_patterns = {
            'EXPERIENCE': [
                r'\b(EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|WORK HISTORY|PROFESSIONAL EXPERIENCE)\b',
            ],
            'EDUCATION': [
                r'\b(EDUCATION|ACADEMIC|DEGREES?|ACADEMIC BACKGROUND)\b',
            ],
            'SKILLS': [
                r'\b(SKILLS|TECHNICAL SKILLS|TECHNOLOGIES|COMPETENCIES|EXPERTISE)\b',
            ],
            'PROJECTS': [
                r'\b(PROJECTS?|PORTFOLIO|NOTABLE PROJECTS?)\b',
            ],
            'CERTIFICATIONS': [
                r'\b(CERTIFICATIONS?|LICENSES?|CREDENTIALS?)\b',
            ],
            'AWARDS': [
                r'\b(AWARDS?|HONORS?|ACHIEVEMENTS?|RECOGNITION)\b',
            ],
            'SUMMARY': [
                r'\b(SUMMARY|PROFESSIONAL SUMMARY|PROFILE|OBJECTIVE)\b',
            ]
        }

        for section_type, patterns in section_patterns.items():
            for pattern in patterns:
                if re.search(pattern, line_clean):
                    return section_type

        return None

    def _split_header_line(self, line: str) -> tuple:
        """
        Smart line splitting for job/project headers
        Returns (main_part, secondary_part)
        """
        # Strategy 1: Split by pipe |
        if '|' in line:
            parts = line.split('|', 1)
            return parts[0].strip(), parts[1].strip() if len(parts) > 1 else ""

        # Strategy 2: Split by 3+ spaces
        if '   ' in line:
            parts = re.split(r'\s{3,}', line, 1)
            return parts[0].strip(), parts[1].strip() if len(parts) > 1 else ""

        # Strategy 3: Check for date at end
        date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*Present|\d{4}\s*[–-]\s*\d{4}|\d{4}\s*[–-]\s*Present'
        date_match = re.search(date_pattern, line, re.IGNORECASE)
        if date_match:
            date_str = date_match.group(0)
            before_date = line[:date_match.start()].strip()
            return before_date, date_str

        # Strategy 4: Check for location at end (City, ST or Remote)
        loc_pattern = r'(.*?)\s+(Remote|Hybrid|On-site|[A-Z][a-z]+,\s*[A-Z]{2})$'
        loc_match = re.search(loc_pattern, line)
        if loc_match:
            return loc_match.group(1).strip(), loc_match.group(2).strip()

        return line.strip(), ""

    async def _parse_with_llm(self, text: str) -> Resume:
        """
        Parse resume using GPT-4 with ROAST framework.
        Handles any resume format accurately.
        """
        import json
        
        # ROAST Framework Prompt for Resume Parsing
        system_prompt = """ROLE: You are an expert resume parser and data extraction specialist with deep knowledge of resume formats, ATS systems, and hiring practices across all industries.

OBJECTIVE: Extract structured resume data from raw text and convert it into a complete, accurate JSON representation that captures:
- Personal information (name, title, contact details)
- Professional summary/objective (ONLY if present)
- Work experience (with all bullets, dates, locations, companies)
- Projects (with descriptions, technologies, URLs) (ONLY if present)
- Skills (categorized by type)
- Education (degrees, institutions, dates, GPAs, honors)
- Certifications, awards (ONLY if present)
- Document structure metadata (section order, which sections exist)

AUDIENCE: Your output will be used by:
- Resume optimization platforms that need accurate data extraction
- ATS systems that require structured data
- Career coaches analyzing resume content
- Job seekers editing their resumes who need their EXACT original structure preserved

STYLE:
- Extract ONLY information that is actually present in the resume
- NEVER add sections that don't exist in the original
- NEVER invent or fabricate content
- Preserve exact wording from bullets and descriptions (word-for-word accuracy)
- Extract contact info (email, phone, LinkedIn, GitHub, website, location)
- Parse dates in any format (convert to consistent format like "Jan 2024" or "2024")
- Categorize skills intelligently (Languages, Frameworks, Tools, etc.)
- Extract project URLs and technologies
- For missing sections, use empty arrays [] or empty strings ""
- Preserve the ORDER of sections as they appear in the original document
- Note which sections are present vs absent

TWO-COLUMN LAYOUT HANDLING:
- Many resumes use two-column layouts where information is spatially separated
- Example: "OPEN SOURCE                                Remote" means company="OPEN SOURCE" and location="Remote"
- Example: "Software Engineer                     Sep 2025 – Nov 2025" means title="Software Engineer" and dates="Sep 2025 – Nov 2025"
- When you see significant spacing (multiple spaces/tabs) between text on the same line, treat them as separate fields
- Company/Location often appear on one line, Title/Dates on the next line
- Parse these spatial relationships correctly into the appropriate JSON fields

TONE: Precise, thorough, and detail-oriented. Extract everything accurately without making assumptions or adding content not in the original."""

        # Create the JSON schema for the response
        schema_instruction = """
Return ONLY valid JSON matching this exact structure:
{
  "header": {
    "name": "string",
    "title": "string",
    "contact": {
      "email": "string",
      "phone": "string",
      "location": "string",
      "linkedin": "string",
      "github": "string",
      "website": "string"
    }
  },
  "summary": "string",
  "experience": [
    {
      "company": "string",
      "title": "string",
      "location": "string",
      "start_date": "string",
      "end_date": "string",
      "current": false,
      "bullets": ["string"],
      "technologies": ["string"]
    }
  ],
  "projects": [
    {
      "name": "string",
      "description": "string",
      "url": "string",
      "start_date": "string",
      "end_date": "string",
      "bullets": ["string"],
      "technologies": ["string"]
    }
  ],
  "skills": {
    "category_name": ["skill1", "skill2"]
  },
  "education": [
    {
      "institution": "string",
      "degree": "string",
      "field": "string",
      "location": "string",
      "start_date": "string",
      "end_date": "string",
      "gpa": "string",
      "honors": ["string"],
      "coursework": ["string"]
    }
  ],
  "certifications": [
    {
      "name": "string",
      "issuer": "string",
      "date": "string",
      "expiry": "string",
      "credential_id": "string",
      "url": "string"
    }
  ],
  "awards": [
    {
      "name": "string",
      "issuer": "string",
      "date": "string",
      "description": "string"
    }
  ]
}

CRITICAL RULES:
1. Extract ONLY sections that actually exist in the resume
2. If a section is missing from the resume, use empty array [] or empty string ""
3. NEVER fabricate or invent sections that aren't in the original
4. Preserve exact wording from bullets - do not paraphrase or improve them
5. Keep the section order as it appears in the original resume"""

        user_prompt = f"""Extract all information from this resume text and return it as structured JSON:

{text[:20000]}

CRITICAL REQUIREMENTS:
1. Extract ALL sections present in the resume (Summary, Experience, Projects, Education, Skills, etc.)
2. For PROJECTS section:
   - Extract EVERY project with its FULL NAME/TITLE
   - If project name appears as a bullet point like "• Built an AI resume optimizer...", treat "Built an AI resume optimizer" as the name
   - Include ALL bullet points under each project (do not truncate or cut off mid-sentence)
   - Preserve complete sentences - if a bullet is "Built an AI web agent that receives natural language tasks and autonomously navigates websites to complete them", include the ENTIRE sentence
3. For EXPERIENCE section:
   - Parse two-column layouts correctly (company/location on one line, title/dates on the next line)
   - Extract ALL bullets completely - do not cut off mid-sentence
4. For EDUCATION section:
   - Extract degree, institution, location, dates
   - Format: "B.S. in Mathematics-Computer Science" at "UNIVERSITY OF CALIFORNIA, SAN DIEGO"
5. Preserve exact wording of ALL bullet points - no truncation, no summarization
6. Extract all contact information from the header
7. DO NOT use special Unicode characters like zero-width spaces in bullets"""

        try:
            # Use AsyncOpenAI for async support
            from openai import AsyncOpenAI
            client = AsyncOpenAI(api_key=self.api_key)
            
            response = await client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[
                    {"role": "system", "content": system_prompt + schema_instruction},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,  # Low temperature for accuracy
                response_format={"type": "json_object"}
            )
            
            result_text = response.choices[0].message.content
            parsed_data = json.loads(result_text)

            # Add structure metadata to track which sections exist
            from app.schemas.resume import DocumentStructure

            structure = DocumentStructure(
                has_summary=bool(parsed_data.get("summary", "").strip()),
                has_projects=len(parsed_data.get("projects", [])) > 0,
                has_certifications=len(parsed_data.get("certifications", [])) > 0,
                has_awards=len(parsed_data.get("awards", [])) > 0,
                sections_present=[],
                section_order=[]
            )

            # Determine section order and presence
            section_map = {
                "header": bool(parsed_data.get("header")),
                "summary": structure.has_summary,
                "experience": len(parsed_data.get("experience", [])) > 0,
                "projects": structure.has_projects,
                "education": len(parsed_data.get("education", [])) > 0,
                "skills": len(parsed_data.get("skills", {})) > 0,
                "certifications": structure.has_certifications,
                "awards": structure.has_awards
            }

            # Typical order for most resumes (can be customized based on detection)
            typical_order = ["header", "summary", "experience", "projects", "education", "skills", "certifications", "awards"]

            for section in typical_order:
                if section_map.get(section, False):
                    structure.sections_present.append(section)
                    structure.section_order.append(section)

            parsed_data["structure"] = structure.dict()

            # Validate and convert to Resume schema
            return Resume(**parsed_data)
            
        except Exception as e:
            print(f"LLM parsing error: {e}")
            import traceback
            traceback.print_exc()
            raise

    async def _parse_basic(self, text: str) -> Resume:
        """GENERALIZABLE resume parser - works with many formats"""
        lines = [line.strip() for line in text.split("\n") if line.strip()]

        print("DEBUG: First 1000 chars of extracted text:")
        print(text[:1000])
        print("=" * 50)

        resume_data = {
            "header": {
                "name": lines[0] if lines else "Unknown",
                "title": "",
                "contact": {}
            },
            "summary": "",
            "experience": [],
            "projects": [],
            "skills": {},
            "education": [],
            "certifications": [],
            "awards": [],
            "flexible_sections": []
        }

        # Extract contact info
        import re
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            resume_data["header"]["contact"]["email"] = email_match.group()

        phone_match = re.search(r'[\+\(]?\d{1,3}[\)\-\.\s]?\d{3}[\-\.\s]?\d{3,4}[\-\.\s]?\d{4}', text)
        if phone_match:
            resume_data["header"]["contact"]["phone"] = phone_match.group()

        linkedin_match = re.search(r'linkedin\.com/in/([\w\-]+)', text, re.IGNORECASE)
        if linkedin_match:
            resume_data["header"]["contact"]["linkedin"] = linkedin_match.group(0)

        github_match = re.search(r'github\.com/([\w\-]+)', text, re.IGNORECASE)
        if github_match:
            resume_data["header"]["contact"]["github"] = github_match.group(0)

        # GENERALIZABLE SUMMARY PARSING
        summary_pattern = r'(SUMMARY|PROFESSIONAL SUMMARY|PROFILE|OBJECTIVE)'
        summary_match = re.search(summary_pattern, text, re.IGNORECASE)
        
        if summary_match:
            summary_start = summary_match.end()
            # Find next section (Experience, Education, Skills, etc.)
            next_section = re.search(r'\n(EXPERIENCE|EDUCATION|SKILLS|PROJECTS|WORK EXPERIENCE)', text[summary_start:], re.IGNORECASE)
            summary_text = text[summary_start:summary_start + next_section.start()] if next_section else text[summary_start:summary_start+500]
            
            # Clean up summary text
            summary_lines = [line.strip() for line in summary_text.split('\n') if line.strip()]
            # Remove section headers that might have been included
            summary_lines = [line for line in summary_lines if not self._detect_section(line)]
            resume_data["summary"] = ' '.join(summary_lines).strip()[:500]  # Limit length

        # GENERALIZABLE EXPERIENCE PARSING with multi-strategy bullet detection
        exp_pattern = r'(EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT|WORK HISTORY|PROFESSIONAL EXPERIENCE)'
        exp_match = re.search(exp_pattern, text, re.IGNORECASE)

        if exp_match:
            exp_start = exp_match.end()
            next_section = re.search(r'\n(EDUCATION|SKILLS|PROJECTS|CERTIFICATIONS)', text[exp_start:], re.IGNORECASE)
            exp_text = text[exp_start:exp_start + next_section.start()] if next_section else text[exp_start:exp_start+6000]

            print(f"DEBUG: Experience section length: {len(exp_text)} chars")

            # Split by double line breaks to separate jobs
            job_blocks = re.split(r'\n\s*\n', exp_text)

            date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*Present|\d{4}\s*[–-]\s*\d{4}|\d{4}\s*[–-]\s*Present'

            for block_idx, block in enumerate(job_blocks):
                if not block.strip():
                    continue

                raw_lines = block.split('\n')
                lines = [l.strip() for l in raw_lines if l.strip()]

                # Use multi-strategy bullet detection
                bullet_lines = []
                for i, raw_line in enumerate(raw_lines):
                    if self._is_bullet_point(raw_line):
                        bullet_lines.append(lines[min(i, len(lines)-1)] if i < len(lines) else raw_line.strip())

                # Also check in stripped lines
                for line in lines:
                    if self._is_bullet_point(line) and line not in bullet_lines:
                        bullet_lines.append(line)

                if not bullet_lines:  # No bullets means not a job entry
                    continue

                # Extract job metadata from lines BEFORE bullets
                header_lines = []
                for l in lines:
                    if self._is_bullet_point(l):
                        break
                    header_lines.append(l)

                if not header_lines:
                    continue

                # Parse header lines - improved logic for separate lines
                company = "Company"
                title = "Position"
                location = ""
                start_date = ""
                end_date = ""

                used_lines = set()  # Track which lines we've used

                # Step 1: Find company (usually all caps or first line)
                for idx, line in enumerate(header_lines):
                    if line.isupper() and len(line) > 2:
                        if not any(month in line.upper() for month in ['JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN', 'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC']):
                            company = line
                            used_lines.add(idx)
                            break

                # If no all-caps company found, use first line
                if company == "Company" and header_lines:
                    main_part, secondary = self._split_header_line(header_lines[0])
                    company = main_part
                    used_lines.add(0)

                # Step 2: Find dates
                for idx, line in enumerate(header_lines):
                    if idx in used_lines:
                        continue
                    date_match = re.search(date_pattern, line, re.IGNORECASE)
                    if date_match:
                        dates_str = date_match.group(0)
                        date_parts = re.split(r'\s*[–-]\s*', dates_str)
                        start_date = date_parts[0].strip() if date_parts else ""
                        end_date = date_parts[1].strip() if len(date_parts) >= 2 else ""
                        used_lines.add(idx)

                        # Check if there's text before the dates on the same line (might be title or location)
                        before_dates = line[:date_match.start()].strip()
                        if before_dates:
                            # Determine if it's a title or location
                            if 'remote' in before_dates.lower() or 'hybrid' in before_dates.lower() or re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}', before_dates):
                                location = before_dates
                            elif not before_dates.isupper():
                                title = before_dates
                        break

                # Step 3: Find location in remaining lines (if not found)
                if not location:
                    for idx, line in enumerate(header_lines):
                        if idx in used_lines:
                            continue
                        # Location keywords
                        if 'remote' in line.lower() or 'hybrid' in line.lower() or 'on-site' in line.lower():
                            location = line
                            used_lines.add(idx)
                            break
                        # City, State pattern (e.g., "New York, NY")
                        if re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}', line):
                            location = line
                            used_lines.add(idx)
                            break

                # Step 4: Find title in remaining lines (if not found)
                if title == "Position":
                    for idx, line in enumerate(header_lines):
                        if idx in used_lines:
                            continue
                        if not line.isupper() and len(line) > 2:
                            if not re.search(date_pattern, line, re.IGNORECASE):
                                # This should be the title
                                title = line
                                used_lines.add(idx)
                                break

                # Clean bullets - remove bullet characters and numbering
                clean_bullets = []
                for bullet in bullet_lines:
                    # Remove bullet chars, numbers, leading dashes/stars
                    clean = re.sub(r'^[\d\.\)\]\}\-\*\•\●\○\▪\▫\■\□\◦\‣\⁃\▸\▹\►\▻]+\s*', '', bullet).strip()
                    if clean and len(clean) > 5:  # Skip very short bullets
                        clean_bullets.append(clean[:500])

                if clean_bullets:
                    job_entry = {
                        "company": company[:100],
                        "title": title[:100],
                        "location": location[:100],
                        "start_date": start_date,
                        "end_date": end_date,
                        "current": "present" in end_date.lower() if end_date else False,
                        "bullets": clean_bullets,
                        "technologies": []
                    }
                    resume_data["experience"].append(job_entry)
                    print(f"DEBUG: Added job: {company} - {title} (Location: {location}) with {len(clean_bullets)} bullets")

        # GENERALIZABLE PROJECTS PARSING with multi-strategy bullet detection
        proj_pattern = r'(PROJECTS?|PORTFOLIO|NOTABLE PROJECTS?)'
        proj_match = re.search(proj_pattern, text, re.IGNORECASE)

        if proj_match:
            proj_start = proj_match.end()
            next_section = re.search(r'\n(EDUCATION|SKILLS|CERTIFICATIONS|AWARDS|EXPERIENCE)', text[proj_start:], re.IGNORECASE)
            proj_text = text[proj_start:proj_start + next_section.start()] if next_section else text[proj_start:proj_start+4000]

            print(f"DEBUG: Projects section length: {len(proj_text)} chars")
            try:
                print(f"DEBUG: Projects section first 500 chars (repr):\n{repr(proj_text[:500])}")
            except Exception as e:
                print(f"DEBUG: Error printing projects text: {e}")

            # Split by paragraphs to separate projects
            proj_blocks = re.split(r'\n\s*\n', proj_text)
            print(f"DEBUG: Found {len(proj_blocks)} project blocks after splitting")

            if len(proj_blocks) == 0:
                print("DEBUG: ERROR - No project blocks found after splitting!")

            for block_idx, block in enumerate(proj_blocks):
                if not block.strip():
                    continue

                raw_lines = block.split('\n')
                lines = [l.strip() for l in raw_lines if l.strip()]

                # Use multi-strategy bullet detection
                bullet_lines = []
                for i, raw_line in enumerate(raw_lines):
                    if self._is_bullet_point(raw_line):
                        bullet_lines.append(lines[min(i, len(lines)-1)] if i < len(lines) else raw_line.strip())

                # Also check in stripped lines
                for line in lines:
                    if self._is_bullet_point(line) and line not in bullet_lines:
                        bullet_lines.append(line)

                print(f"DEBUG: Project block {block_idx}: {len(lines)} lines, {len(bullet_lines)} bullets")
                if not lines:
                    print(f"DEBUG: Skipping block {block_idx} - no lines")
                    continue

                # Projects might not have bullets - could just be name + description
                # Extract project header - if ALL lines are bullets, treat first line as header
                header_lines = []
                for l in lines:
                    if self._is_bullet_point(l):
                        break
                    header_lines.append(l)

                # If no header found (all bullets), use first line as project name
                if not header_lines and lines:
                    header_lines = [lines[0]]
                    # Remove first line from bullet_lines since it's the project name
                    if bullet_lines and bullet_lines[0] == lines[0]:
                        bullet_lines = bullet_lines[1:]
                    print(f"DEBUG: Project block {block_idx} - using first line as header: {lines[0][:50]}")
                elif not header_lines:
                    print(f"DEBUG: Skipping block {block_idx} - no header lines and no lines at all")
                    continue

                # First line is usually: "Project Name | Tech1, Tech2, Tech3" or "Project Name (url)"
                first_line = header_lines[0]
                project_name = first_line
                technologies = []
                description = ""
                url = ""
                start_date = ""
                end_date = ""

                # Try smart line splitting for name and tech stack
                main_part, tech_part = self._split_header_line(first_line)
                project_name = main_part

                # Extract technologies from tech_part
                if tech_part:
                    # Check if it's a date or tech stack
                    date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|\d{4}'
                    if re.search(date_pattern, tech_part, re.IGNORECASE):
                        # It's a date range
                        date_match = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*Present|\d{4}\s*[–-]\s*\d{4}|\d{4}\s*[–-]\s*Present', tech_part, re.IGNORECASE)
                        if date_match:
                            date_parts = re.split(r'\s*[–-]\s*', date_match.group(0))
                            start_date = date_parts[0].strip() if date_parts else ""
                            end_date = date_parts[1].strip() if len(date_parts) >= 2 else ""
                    else:
                        # It's a tech stack
                        technologies = [t.strip() for t in re.split(r'[,;]', tech_part) if t.strip()]

                # Check for URL in parentheses
                url_match = re.search(r'\(([^)]+)\)', project_name)
                if url_match:
                    url = url_match.group(1)
                    project_name = project_name.replace(url_match.group(0), '').strip()

                # Check for technologies in second line if not found
                if not technologies and len(header_lines) > 1:
                    second_line = header_lines[1]
                    # If second line looks like a tech list (has commas or common tech keywords)
                    if ',' in second_line or any(tech in second_line.lower() for tech in ['python', 'java', 'react', 'node', 'javascript', 'typescript', 'c++', 'go', 'rust']):
                        technologies = [t.strip() for t in re.split(r'[,;]', second_line) if t.strip()]
                    else:
                        # It's a description
                        description = second_line

                # Third line might be description if not already set
                if not description and len(header_lines) > 2:
                    description = header_lines[2]

                # Clean bullets - remove bullet characters and numbering
                clean_bullets = []
                for bullet in bullet_lines:
                    clean = re.sub(r'^[\d\.\)\]\}\-\*\•\●\○\▪\▫\■\□\◦\‣\⁃\▸\▹\►\▻]+\s*', '', bullet).strip()
                    if clean and len(clean) > 5:  # Skip very short bullets
                        clean_bullets.append(clean[:500])

                # Add project if we have meaningful content
                print(f"DEBUG: Project block {block_idx} evaluation - bullets: {len(clean_bullets)}, description: {bool(description)}, header_lines: {len(header_lines)}, name: {project_name[:50] if project_name else 'NONE'}")
                if clean_bullets or description or len(header_lines) > 1:
                    project_entry = {
                        "name": project_name[:100],
                        "description": description[:300],
                        "url": url[:200],
                        "start_date": start_date,
                        "end_date": end_date,
                        "bullets": clean_bullets,
                        "technologies": technologies[:20]
                    }
                    resume_data["projects"].append(project_entry)
                    print(f"DEBUG: Added project: {project_name} with {len(clean_bullets)} bullets, {len(technologies)} techs")
                else:
                    print(f"DEBUG: Skipping project block {block_idx} - no bullets, no description, only 1 header line")

        # GENERALIZABLE SKILLS PARSING
        skills_pattern = r'(SKILLS|TECHNICAL SKILLS|TECHNOLOGIES|COMPETENCIES|EXPERTISE)'
        skills_match = re.search(skills_pattern, text, re.IGNORECASE)
        if skills_match:
            skills_start = skills_match.end()
            next_section = re.search(r'\n(EXPERIENCE|EDUCATION|PROJECTS|CERTIFICATIONS|AWARDS)', text[skills_start:], re.IGNORECASE)
            skills_text = text[skills_start:skills_start + next_section.start()] if next_section else text[skills_start:skills_start+800]

            # Parse by category if possible (most common format)
            # Example: "Languages: Python, Java, C++"
            category_pattern = r'([A-Za-z\s&/]+?):\s*([^\n]+)'
            categories = re.findall(category_pattern, skills_text)

            if categories:
                for cat_name, cat_skills in categories:
                    # Split by commas, semicolons, or bullets
                    skills_list = [s.strip() for s in re.split(r'[,;•●]', cat_skills) if s.strip()]
                    if skills_list:
                        resume_data["skills"][cat_name.strip()] = skills_list
            else:
                # Fallback: Look for bullet points or comma-separated lists
                skill_lines = [l.strip() for l in skills_text.split('\n') if l.strip()]
                all_skills = []
                for line in skill_lines:
                    # Check if it's a bullet point
                    if self._is_bullet_point(line):
                        clean = re.sub(r'^[\d\.\)\]\}\-\*\•\●\○\▪\▫\■\□\◦\‣\⁃\▸\▹\►\▻]+\s*', '', line).strip()
                        if clean:
                            all_skills.append(clean)
                    # Check if it has commas (comma-separated list)
                    elif ',' in line:
                        skills = [s.strip() for s in line.split(',') if s.strip()]
                        all_skills.extend(skills)
                    # Single skill per line
                    elif len(line) > 2 and not self._detect_section(line):
                        all_skills.append(line)

                if all_skills:
                    resume_data["skills"]["Technical"] = all_skills[:30]

        # GENERALIZABLE EDUCATION PARSING
        edu_pattern = r'(EDUCATION|ACADEMIC|DEGREES?|ACADEMIC BACKGROUND)'
        edu_match = re.search(edu_pattern, text, re.IGNORECASE)
        if edu_match:
            edu_start = edu_match.end()
            next_section = re.search(r'\n(EXPERIENCE|SKILLS|PROJECTS|CERTIFICATIONS)', text[edu_start:], re.IGNORECASE)
            edu_text = text[edu_start:edu_start + next_section.start()] if next_section else text[edu_start:edu_start+800]

            # Split by paragraphs to separate multiple degrees
            edu_blocks = re.split(r'\n\s*\n', edu_text)

            for block in edu_blocks:
                if not block.strip():
                    continue

                edu_lines = [l.strip() for l in block.split('\n') if l.strip() and len(l.strip()) > 2]
                if not edu_lines:
                    continue

                institution = ""
                degree = ""
                field = ""
                location = ""
                start_date = ""
                end_date = ""
                gpa = ""
                used_line_indices = set()

                # Find institution (usually all caps or first line)
                for idx, line in enumerate(edu_lines):
                    if line.isupper() and len(line) > 3:
                        # Make sure it's not a date or degree line
                        if not any(keyword in line.lower() for keyword in ['bachelor', 'master', 'phd', 'b.s.', 'm.s.', 'b.a.', 'm.a.']) and not re.search(r'\d{4}', line):
                            institution = line
                            used_line_indices.add(idx)
                            break

                if not institution and edu_lines:
                    # First line is likely institution if it's not a degree line
                    if not any(keyword in edu_lines[0].lower() for keyword in ['bachelor', 'master', 'phd', 'b.s.', 'm.s.', 'b.a.', 'm.a.']):
                        institution = edu_lines[0]
                        used_line_indices.add(0)

                # Find degree (often has "Bachelor", "Master", "PhD", "B.S.", "M.S.", etc.)
                degree_keywords = ['bachelor', 'master', 'phd', 'ph.d', 'b.s.', 'm.s.', 'b.a.', 'm.a.', 'associate', 'doctorate', 'diploma']
                for idx, line in enumerate(edu_lines):
                    if idx in used_line_indices:
                        continue
                    if any(keyword in line.lower() for keyword in degree_keywords):
                        # Check if field is in same line (e.g., "Bachelor of Science in Computer Science")
                        if ' in ' in line.lower():
                            # Split on ' in ' (case-insensitive)
                            parts = re.split(r'\s+in\s+', line, maxsplit=1, flags=re.IGNORECASE)
                            if len(parts) == 2:
                                degree = parts[0].strip()  # Just the degree part
                                field = parts[1].strip()   # The field (preserve case)
                            else:
                                degree = line
                        else:
                            degree = line
                        used_line_indices.add(idx)
                        break

                # Find dates
                date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s*[–-]\s*Present|\d{4}\s*[–-]\s*\d{4}|\d{4}\s*[–-]\s*Present|\b(19|20)\d{2}\b'
                for line in edu_lines:
                    date_match = re.search(date_pattern, line, re.IGNORECASE)
                    if date_match:
                        dates_str = date_match.group(0)
                        # Check if it's a range or single year
                        if '–' in dates_str or '-' in dates_str:
                            date_parts = re.split(r'\s*[–-]\s*', dates_str)
                            start_date = date_parts[0].strip() if date_parts else ""
                            end_date = date_parts[1].strip() if len(date_parts) >= 2 else ""
                        else:
                            # Single year - probably graduation year
                            end_date = dates_str.strip()
                        break

                # Find GPA
                gpa_match = re.search(r'GPA:?\s*([\d\.]+)', edu_text, re.IGNORECASE)
                if gpa_match:
                    gpa = gpa_match.group(1)

                # Find location
                for line in edu_lines:
                    if re.search(r'[A-Z][a-z]+,\s*[A-Z]{2}', line) and not degree:
                        location = line

                resume_data["education"].append({
                    "institution": institution[:100],
                    "degree": degree[:100],
                    "field": field[:100],
                    "location": location[:100],
                    "start_date": start_date,
                    "end_date": end_date,
                    "gpa": gpa,
                    "honors": [],
                    "coursework": []
                })

        print(f"\nDEBUG: Final parsing results:")
        print(f"  - Experience items: {len(resume_data['experience'])}")
        print(f"  - Projects: {len(resume_data['projects'])}")
        print(f"  - Skills categories: {len(resume_data['skills'])}")
        print(f"  - Education items: {len(resume_data['education'])}")

        return Resume(**resume_data)


# Create service instance for backward compatibility
class ResumeParserService:
    """Main parser service using LLM"""

    def __init__(self):
        self.llm_parser = LLMResumeParser()

    async def parse_file(self, file_path: Path, filename: str) -> Resume:
        return await self.llm_parser.parse_file(file_path, filename)
